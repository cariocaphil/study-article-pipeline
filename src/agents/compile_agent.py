"""
Compile agent.
Takes a validated PipelineOutput and produces a printable .docx file.
"""

import logging

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from src.schemas.article import PipelineOutput

logger = logging.getLogger(__name__)


def compile_document(output: PipelineOutput, output_path: str) -> str:
    """
    Generates a .docx file from a PipelineOutput.
    Returns the path to the generated file.
    """

    doc = Document()

    # ── Page margins — generous for hand annotation ───────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(4)  # wider right margin for your pen notes

    # ── Document title ────────────────────────────────────────────────────────
    title = doc.add_heading(f'Study Article Collection regarding "{output.topic}"', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Metadata line ─────────────────────────────────────────────────────────
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Source language: {output.source_language.capitalize()}  |  "
        f"Translations: {output.translation_language.capitalize()}  |  "
        f"Level: {output.user_level.value}"
    ).italic = True

    doc.add_paragraph()  # spacer

    # ── Articles ──────────────────────────────────────────────────────────────
    for i, article in enumerate(output.articles, start=1):
        # Article heading
        doc.add_heading(f"Article {i}", level=2)

        # Article metadata
        p = doc.add_paragraph()
        p.add_run("Title: ").bold = True
        p.add_run(article.title)

        p = doc.add_paragraph()
        p.add_run("Author: ").bold = True
        p.add_run(article.author if article.author else "Unknown")

        p = doc.add_paragraph()
        p.add_run("Source: ").bold = True
        p.add_run(article.source_name)

        p = doc.add_paragraph()
        p.add_run("Link: ").bold = True
        p.add_run(article.url)

        doc.add_paragraph()  # spacer

        # Full article text
        doc.add_heading("Article Text", level=3)
        for paragraph in article.full_text.split("\n"):
            paragraph = paragraph.strip()
            if paragraph:
                doc.add_paragraph(paragraph)

        doc.add_paragraph()  # spacer

        # Vocabulary and expressions list
        doc.add_heading("Vocabulary & Expressions", level=3)

        if not article.phrases:
            doc.add_paragraph("No phrases extracted above the specified level.")
        else:
            for phrase in article.phrases:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(
                    f"[{phrase.category.value}] [{phrase.estimated_level.value}] "
                ).bold = True
                p.add_run(f"{phrase.phrase}").bold = True
                p.add_run(f" — {phrase.translation}")
                # Sentence context on the next line, indented
                context_p = doc.add_paragraph(style="List Bullet")
                context_p.add_run(f'"{phrase.sentence_context}"').italic = True
                context_p.paragraph_format.left_indent = Cm(1)

        # Page break between articles (except after the last one)
        if i < len(output.articles):
            doc.add_page_break()

    doc.save(output_path)
    logger.info("Document saved to %s", output_path)
    return output_path


# ── Manual test ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    from src.schemas.article import Article, CEFRLevel, ExtractedPhrase, PhraseCategory

    sample = PipelineOutput(
        topic="Entroncamento",
        source_language="portuguese",
        translation_language="german",
        user_level=CEFRLevel.C1,
        articles=[
            Article(
                title="Entroncamento, a Crítica | Pedro Cabeleira revela a sensibilidade na violência",
                author="Matilde Sousa",
                url="https://www.magazine-hd.com/apps/wp/entroncamento-critica-filme-pedro-cabeleira-ana-vilaca/",
                source_name="magazine-hd.com",
                full_text=(
                    "Em «Entroncamento» acompanhamos Laura, que foge de um passado turbulento, "
                    "refugiando-se nesta cidade do distrito de Santarém para recomeçar a sua vida. "
                    "Contudo, apesar de tentar encontrar um emprego honesto e uma vida melhor, "
                    "começa a entrar em pequenos esquemas e crimes, motivados por familiares e "
                    "conhecidos que a envolvem numa teia de cumplicidades difícil de escapar."
                ),
                phrases=[
                    ExtractedPhrase(
                        phrase="entrar em pequenos esquemas",
                        sentence_context="começa a entrar em pequenos esquemas e crimes, motivados por familiares e conhecidos",
                        translation="in kleine Machenschaften verwickelt werden",
                        category=PhraseCategory.idiom,
                        estimated_level=CEFRLevel.C1,
                    ),
                    ExtractedPhrase(
                        phrase="teia de cumplicidades",
                        sentence_context="que a envolvem numa teia de cumplicidades difícil de escapar",
                        translation="Netz der Komplizenschaft",
                        category=PhraseCategory.idiom,
                        estimated_level=CEFRLevel.C1,
                    ),
                ],
            )
        ],
    )

    compile_document(sample, "output/test_output.docx")
